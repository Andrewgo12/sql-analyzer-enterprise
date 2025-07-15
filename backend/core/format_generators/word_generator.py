"""
Word Document Generator

Generates professional Word documents with analysis results.
"""

import time
import io
from typing import Dict, Any
from .base_generator import BaseFormatGenerator, GenerationContext, GenerationResult

# Optional import for Word document generation
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.shared import OxmlElement, qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    # Create dummy classes for when python-docx is not available
    class Document:
        def __init__(self, *args, **kwargs):
            pass
        def add_heading(self, *args, **kwargs):
            return self
        def add_paragraph(self, *args, **kwargs):
            return self
        def add_table(self, *args, **kwargs):
            return self
        def save(self, *args, **kwargs):
            pass

    class Inches:
        def __init__(self, *args, **kwargs):
            pass

    class Pt:
        def __init__(self, *args, **kwargs):
            pass

    WD_ALIGN_PARAGRAPH = type('WD_ALIGN_PARAGRAPH', (), {'CENTER': 1, 'LEFT': 0})
    WD_STYLE_TYPE = type('WD_STYLE_TYPE', (), {'PARAGRAPH': 1})


class WordDocumentGenerator(BaseFormatGenerator):
    """Generator for professional Word documents."""

    @property
    def format_name(self) -> str:
        return "Documento Word"

    @property
    def file_extension(self) -> str:
        return ".docx"

    @property
    def mime_type(self) -> str:
        return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    @property
    def is_binary(self) -> bool:
        return True

    def generate(self, context: GenerationContext) -> GenerationResult:
        """Generate professional Word document."""
        start_time = time.time()

        try:
            if not DOCX_AVAILABLE:
                # Create a simple text-based document if python-docx is not available
                return self._generate_text_fallback(context, start_time)

            self.validate_context(context)

            # Generate Word document
            doc_content = self._generate_word_document(context)

            generation_time = time.time() - start_time

            return self.create_generation_result(
                doc_content, context, generation_time,
                metadata={
                    "pages_estimated": self._estimate_page_count(context),
                    "sections": 5,
                    "format": "DOCX"
                }
            )

        except Exception as e:
            return self.handle_generation_error(e, context)

    def _generate_word_document(self, context: GenerationContext) -> bytes:
        """Generate the complete Word document."""
        doc = Document()

        # Set document properties
        doc.core_properties.title = f"Análisis SQL - {context.original_filename}"
        doc.core_properties.author = "SQL Analyzer Enterprise"
        doc.core_properties.subject = "Reporte de Análisis SQL"

        # Get template variables
        template_vars = self.get_template_variables(context)

        # Add title page
        self._add_title_page(doc, template_vars)

        # Add page break
        doc.add_page_break()

        # Add executive summary
        self._add_executive_summary(doc, template_vars)

        # Add errors section
        self._add_errors_section(doc, template_vars)

        # Add recommendations
        self._add_recommendations_section(doc, template_vars)

        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        doc_content = buffer.getvalue()
        buffer.close()

        return doc_content

    def _add_title_page(self, doc: Document, template_vars: Dict[str, Any]):
        """Add title page to the document."""
        # Title
        title = doc.add_heading('Reporte de Análisis SQL', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Subtitle
        subtitle = doc.add_heading(f'Análisis de {template_vars["original_filename"]}', 1)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add some spacing
        doc.add_paragraph()
        doc.add_paragraph()

        # Summary table
        summary = template_vars['summary']
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'

        # Header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Métrica'
        hdr_cells[1].text = 'Valor'

        # Data rows
        data = [
            ('Archivo Analizado', template_vars['original_filename']),
            ('Fecha de Análisis', template_vars['analysis_timestamp'].strftime('%d/%m/%Y %H:%M:%S')),
            ('Total de Errores', str(summary['total_errors'])),
            ('Puntuación de Calidad', f"{summary['quality_score']}%"),
            ('Líneas Analizadas', str(summary['lines_analyzed'])),
            ('Correcciones Sugeridas', str(summary['fixes_suggested']))
        ]

        for metric, value in data:
            row_cells = table.add_row().cells
            row_cells[0].text = metric
            row_cells[1].text = value

        # Footer
        doc.add_paragraph()
        footer = doc.add_paragraph('Generado por SQL Analyzer Enterprise')
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _add_executive_summary(self, doc: Document, template_vars: Dict[str, Any]):
        """Add executive summary section."""
        doc.add_heading('Resumen Ejecutivo', 1)

        summary = template_vars['summary']

        # Overview paragraph
        overview = doc.add_paragraph()
        overview.add_run(f"El análisis del archivo ").bold = False
        overview.add_run(f"{template_vars['original_filename']}").bold = True
        overview.add_run(f" ha identificado ").bold = False
        overview.add_run(f"{summary['total_errors']} errores").bold = True
        overview.add_run(f" en total, con una puntuación de calidad de ").bold = False
        overview.add_run(f"{summary['quality_score']}%").bold = True
        overview.add_run(f". Se han analizado ").bold = False
        overview.add_run(f"{summary['lines_analyzed']} líneas").bold = True
        overview.add_run(f" de código SQL.").bold = False

        # Severity breakdown
        doc.add_heading('Distribución de Errores por Severidad', 2)

        severity_table = doc.add_table(rows=1, cols=3)
        severity_table.style = 'Table Grid'

        # Header
        hdr_cells = severity_table.rows[0].cells
        hdr_cells[0].text = 'Severidad'
        hdr_cells[1].text = 'Cantidad'
        hdr_cells[2].text = 'Porcentaje'

        # Data
        severity_data = [
            ('Críticos', summary['critical_errors']),
            ('Altos', summary['high_errors']),
            ('Medios', summary['medium_errors']),
            ('Bajos', summary['low_errors'])
        ]

        for severity, count in severity_data:
            row_cells = severity_table.add_row().cells
            row_cells[0].text = severity
            row_cells[1].text = str(count)
            percentage = (count / max(1, summary['total_errors'])) * 100
            row_cells[2].text = f"{percentage:.1f}%"

    def _add_errors_section(self, doc: Document, template_vars: Dict[str, Any]):
        """Add errors section."""
        doc.add_heading('Errores Detectados', 1)

        errors = template_vars['errors'][:20]  # Limit to first 20 errors

        for i, error in enumerate(errors, 1):
            # Error header
            error_heading = doc.add_heading(f"{i}. {error['title']}", 3)

            # Error details
            details = doc.add_paragraph()
            details.add_run("Severidad: ").bold = True
            details.add_run(f"{error['severity_display']}\n")
            details.add_run("Línea: ").bold = True
            details.add_run(f"{error['line']}\n")
            details.add_run("Descripción: ").bold = True
            details.add_run(f"{error['description']}")

            # Fixes if available
            if error['fixes']:
                fixes_para = doc.add_paragraph()
                fixes_para.add_run("Correcciones Sugeridas:").bold = True

                for fix in error['fixes'][:2]:  # Limit to 2 fixes
                    fix_para = doc.add_paragraph(style='List Bullet')
                    fix_para.add_run(f"{fix.get('description', '')} ")
                    fix_para.add_run(f"(Confianza: {fix.get('confidence', 0)*100:.1f}%)").italic = True

        if len(template_vars['errors']) > 20:
            doc.add_paragraph(f"... y {len(template_vars['errors']) - 20} errores adicionales")

    def _add_recommendations_section(self, doc: Document, template_vars: Dict[str, Any]):
        """Add recommendations section."""
        doc.add_heading('Recomendaciones', 1)

        summary = template_vars['summary']
        recommendations = []

        if summary['critical_errors'] > 0:
            recommendations.append("🚨 URGENTE: Corrija los errores críticos antes de ejecutar este código en producción")

        if summary['high_errors'] > 0:
            recommendations.append("⚠️ IMPORTANTE: Revise y corrija los errores de alta prioridad")

        if summary['quality_score'] < 70:
            recommendations.append("📈 MEJORA: La puntuación de calidad es baja, considere refactorizar el código")

        if summary['fixes_suggested'] > 0:
            recommendations.append(f"🔧 AUTOMATIZACIÓN: {summary['fixes_suggested']} correcciones pueden aplicarse automáticamente")

        recommendations.extend([
            "📚 DOCUMENTACIÓN: Revise la documentación completa para mejores prácticas",
            "🔍 REVISIÓN: Implemente revisiones de código regulares",
            "🧪 TESTING: Pruebe el código en un entorno de desarrollo antes de producción"
        ])

        for recommendation in recommendations:
            doc.add_paragraph(recommendation, style='List Bullet')

    def _generate_text_fallback(self, context: GenerationContext, start_time: float) -> GenerationResult:
        """Generate a text-based fallback when python-docx is not available."""
        template_vars = self.get_template_variables(context)
        summary = template_vars['summary']

        content = f"""REPORTE DE ANÁLISIS SQL
========================

Archivo: {template_vars['original_filename']}
Fecha: {template_vars['analysis_timestamp'].strftime('%d/%m/%Y %H:%M:%S')}

RESUMEN EJECUTIVO
=================
Total de errores: {summary['total_errors']}
Puntuación de calidad: {summary['quality_score']}%
Líneas analizadas: {summary['lines_analyzed']}
Correcciones sugeridas: {summary['fixes_suggested']}

DISTRIBUCIÓN DE ERRORES
========================
Críticos: {summary['critical_errors']}
Altos: {summary['high_errors']}
Medios: {summary['medium_errors']}
Bajos: {summary['low_errors']}

RECOMENDACIONES
===============
- Revise los errores críticos inmediatamente
- Implemente las correcciones sugeridas
- Mejore la calidad general del código

Generado por SQL Analyzer Enterprise
"""

        generation_time = time.time() - start_time

        return self.create_generation_result(
            content.encode('utf-8'), context, generation_time,
            metadata={"fallback": True, "format": "text"}
        )

    def _estimate_page_count(self, context: GenerationContext) -> int:
        """Estimate the number of pages in the document."""
        errors_count = len(context.analysis_result.get("errors", []))

        # Base pages: title, summary, recommendations
        base_pages = 3

        # Estimate pages for errors (approximately 3 errors per page)
        error_pages = (errors_count + 2) // 3

        return base_pages + error_pages
